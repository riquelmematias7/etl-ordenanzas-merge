import re, math, csv, datetime
from pathlib import Path
from docx import Document

# === CONFIGURACIÓN DE CARPETA ===
CARPETA = Path(".")  # Carpeta donde están los archivos DOCX
# ================================

# 🧮 Pedir el porcentaje de aumento
while True:
    try:
        PORCENTAJE_AUMENTO = float(input("📈 Ingresá el porcentaje de aumento (ej: 8.61 para 8.61%): ").strip())
        break
    except ValueError:
        print("❌ Por favor, ingresá un número válido (ejemplo: 10 o 8.61).")

# === DETECCIÓN DE NÚMEROS ===
# Captura cualquier valor tipo $ 2312.21 / $2.732,00 / $1,964.00 / $44,31
pattern = re.compile(r"(\$?\s*)(\d+(?:[.,]\d{1,3})*(?:[.,]\d{2})?)")
factor = 1 + PORCENTAJE_AUMENTO / 100

# === CONVERSIÓN FLEXIBLE ===
def to_float_local(s: str) -> float:
    """
    Convierte cualquier formato numérico (mezclado o incorrecto) a float.
    Ejemplos válidos:
      1.234,56 | 1,234.56 | 2312.21 | 2.732.00 | 44,31 | 205.00
    """
    s = s.strip()
    # eliminar símbolos de moneda, espacios o guiones
    s = re.sub(r"[^\d,\.]", "", s)

    # si tiene más de un punto, el último es decimal
    if s.count(".") > 1:
        parts = s.split(".")
        s = "".join(parts[:-1]) + "." + parts[-1]
    elif s.count(",") > 1:
        parts = s.split(",")
        s = "".join(parts[:-1]) + "," + parts[-1]

    # si hay coma y punto
    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        s = s.replace(".", "").replace(",", ".")
    # ya en formato con punto decimal
    try:
        return float(s)
    except Exception:
        raise ValueError(f"No se pudo convertir el número: {s}")

# === FORMATEO CORRECTO ===
def format_local_int_with_cents(n: int) -> str:
    """Devuelve número con miles por punto y ',00' fijo."""
    return f"{n:,}".replace(",", ".") + ",00"

def gather_runs(paragraph):
    text, spans, hl, pos = "", [], [], 0
    for run in paragraph.runs:
        t = run.text or ""
        start, end = pos, pos + len(t)
        text += t
        spans.append((start, end))
        hl.append(getattr(run.font, "highlight_color", None) is not None)
        pos = end
    return text, spans, hl

def is_highlighted(span, spans, highlights):
    mstart, mend = span
    for (rs, re), h in zip(spans, highlights):
        if h and not (re <= mstart or rs >= mend):
            return True
    return False

def procesar_documento(path: Path, log_file):
    doc = Document(str(path))
    audit = [("ubicacion", "original", "aumentado_antes_redondeo", "nuevo")]
    actualizados = 0

    print(f"\n📄 Procesando: {path.name}")
    log_file.write(f"\n📄 Procesando: {path.name}\n")

    def procesar_parrafo(p, etiqueta):
        nonlocal actualizados
        full, spans, highlights = gather_runs(p)
        if not full:
            return
        nuevo = full
        hubo = False

        for m in sorted(pattern.finditer(full), key=lambda x: x.start(), reverse=True):
            g1, g2 = m.group(1), m.group(2)
            if not is_highlighted((m.start(), m.end()), spans, highlights):
                continue
            try:
                val = to_float_local(g2)
            except Exception:
                continue
            inc = val * factor
            nuevo_val = math.floor(inc + 1e-9)
            nuevo_txt = f"{g1.strip()} {format_local_int_with_cents(nuevo_val)}"

            msg = f"💡 {g1}{g2} → +{PORCENTAJE_AUMENTO}% = {inc:,.2f} → redondeado = {nuevo_txt}"
            print(msg)
            log_file.write(msg + "\n")

            nuevo = nuevo[:m.start()] + nuevo_txt + nuevo[m.end():]
            audit.append((etiqueta, f"{g1}{g2}", f"{inc:.2f}", nuevo_txt))
            hubo = True
            actualizados += 1

        if hubo and p.runs:
            p.runs[0].text = nuevo
            for r in p.runs[1:]:
                r.text = ""

    # Párrafos y tablas
    for i, p in enumerate(doc.paragraphs):
        procesar_parrafo(p, f"parrafo#{i+1}")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    procesar_parrafo(p, f"tabla#{ti+1}:fila#{ri+1}:col#{ci+1}:p#{pi+1}")

    # Guardar resultados
    out_docx = path.with_name(path.stem + f" - Aumento {PORCENTAJE_AUMENTO:.2f}%.docx")
    doc.save(out_docx)
    audit_csv = path.with_name(path.stem + f" - Auditoria {PORCENTAJE_AUMENTO:.2f}%.csv")
    with open(audit_csv, "w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerows(audit)

    resumen = f"✅ {actualizados} valores corregidos y actualizados en {path.name}."
    print(resumen)
    log_file.write(resumen + "\n\n")

    return actualizados

# === PROCESAR TODOS LOS DOCX ===
timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
log_path = CARPETA / f"actualizacion_log_{timestamp}.txt"
total = 0

with open(log_path, "w", encoding="utf-8") as log:
    log.write(f"=== INICIO DEL PROCESO ({timestamp}) ===\n")
    log.write(f"Porcentaje de aumento: {PORCENTAJE_AUMENTO}%\n\n")
    for file in CARPETA.glob("*.docx"):
        if "Aumento" in file.stem:
            continue
        total += procesar_documento(file, log)
    log.write(f"\n📊 Proceso completado. Total de valores corregidos y actualizados: {total}\n")

print(f"\n📊 Proceso completado. Total de valores corregidos y actualizados: {total}")
print(f"🗒️ Log guardado en: {log_path}")